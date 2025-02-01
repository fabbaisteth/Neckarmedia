import glob
import PyPDF2
import docx
import os
import ollama
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.vectorstores.utils import filter_complex_metadata
from dotenv import load_dotenv
from langdetect import detect
import spacy
from transformers import pipeline
import shutil

load_dotenv()

docx_files = glob.glob("docs/*.docx")
API_KEY = os.getenv("GOOGLE_DRIVE_API")
FOLDER_ID = "1af9TUTNrBSkaoHZrSyqYWSTYqk0UiER3"

nlp = spacy.load("de_core_news_md")  # German model
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")

SERVICE_CATEGORIES = [
    "Digital Analytics", "SEO", "SEA", "PLA", "CRO", "Content Marketing", "Social Media Marketing"
]

chroma_db_path = "./chroma_db"
if os.path.exists(chroma_db_path):
    shutil.rmtree(chroma_db_path)  # Deletes the existing database

print("ChromaDB reset. Reinitializing...")

# Reinitialize ChromaDB collection
embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
db = Chroma(collection_name="neckarmedia", embedding_function=embedding, persist_directory=chroma_db_path)

def extract_docx_text(docx_path):
    d = docx.Document(docx_path)
    return "\n".join([p.text for p in d.paragraphs])

def extract_keywords_llm(text, num_keywords=10):
    """Uses the local Mistral model via Ollama to generate relevant keywords."""
    prompt = f"Extract {num_keywords} relevant keywords summarizing this text. Output only keywords, separated by commas.\n\n{text[:1000]}"  # Limit input to 1000 characters

    response = ollama.chat(model="mistral", messages=[{"role": "user", "content": prompt}])

    if "message" in response and "content" in response["message"]:
        return response["message"]["content"].strip()
    
    return "Unknown"

def annotate_metadata(text, file_path):
    metadata = {
        "url": file_path,
        "title": os.path.basename(file_path),
        "category": "",
        "language": str(detect(text)),
        "keywords": extract_keywords_llm(text, num_keywords=10),
        "entities": ""
    }
    
    doc = nlp(text)
    metadata["entities"] = ", ".join([ent.text for ent in doc.ents])  # Convert list to string
    
    try:
        prediction_text = text[:512] if len(text) > 512 else text  # Truncate input
        prediction = classifier(prediction_text, candidate_labels=SERVICE_CATEGORIES)
        metadata["category"] = str(prediction["labels"][0]) if "labels" in prediction else "Unknown"
    except Exception as e:
        print(f"⚠️ Classification failed for {file_path}: {e}")
        metadata["category"] = "Unknown"
    
    return metadata

splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

# Process PDF and DOCX files
for file_path in docx_files:
    raw_text = extract_docx_text(file_path)
    docs = splitter.create_documents([raw_text])
    
    for d in docs:
        metadata = annotate_metadata(d.page_content, file_path)
        d.metadata = metadata
        print(f"Adding document: {d.metadata['title']}")
    db.add_texts([d.page_content for d in docs], [d.metadata for d in docs])

# Process Crawled Web Pages
FILENAME = "neckarmedia_crawl.json"
with open(FILENAME, "r", encoding="utf-8") as f:
    crawled_pages = json.load(f)

documents = []
for page in crawled_pages:
    doc = Document(
        page_content=page["content"],
        metadata={
            "url": page["url"],
            "title": page["title"]
        }
    )
    print(f"Adding document: {doc.metadata['title']}")
    documents.append(doc)

chunked_docs = splitter.split_documents(documents)
print(f"🔄 Processing {len(chunked_docs)} document chunks...")
for i, d in enumerate(chunked_docs):
    metadata = annotate_metadata(d.page_content, d.metadata["url"])
    if i % 10 == 0:
        print(f"⏳ {i}/{len(chunked_docs)} chunks processed")
    d.metadata = metadata
print(f"Chunking done Checkpoint")
db.add_texts([d.page_content for d in chunked_docs], [d.metadata for d in chunked_docs])
print("✅ All documents (crawl + PDFs + DOCX) successfully ingested into ChromaDB! 🎉")

db.persist()

print("Neckarmedia documents and crawled pages ingested into Chroma!")